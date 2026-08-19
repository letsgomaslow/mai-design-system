#!/usr/bin/env python3
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "artifacts" / "docx"
LOGO = ROOT / "artifacts" / "generated" / "maslow-mark-ink.png"

NAVY = "192332"
BAND = "121D35"
TEXT = "333333"
MUTED = "666666"
PURPLE = "654C8F"
LINK = "9D4B8E"
PINK = "EE7BB3"
TEAL = "73C1AE"
OFF_WHITE = "F6F7F9"
LINE = "E1E1E1"
WHITE = "FFFFFF"


def set_run(run, font="Manrope", size=10.5, color=TEXT, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_style_font(style, font, size, color, bold=False):
    style.font.name = font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def define_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Manrope", 10.5, TEXT)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Manrope", 16, NAVY, True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Manrope", 12.5, NAVY, True)
    h2.paragraph_format.space_before = Pt(11)
    h2.paragraph_format.space_after = Pt(5)

    title = doc.styles.add_style("Maslow Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, "Manrope", 26, NAVY, True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles.add_style("Maslow Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, "Manrope", 12, MUTED)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.line_spacing = 1.2

    eyebrow = doc.styles.add_style("Maslow Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(eyebrow, "IBM Plex Mono", 8, PURPLE, True)
    eyebrow.paragraph_format.space_after = Pt(5)

    evidence = doc.styles.add_style("Maslow Evidence", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(evidence, "IBM Plex Mono", 8, PURPLE, True)
    evidence.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths: Iterable[float], indent_dxa=0):
    widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    if layout.getparent() is None:
        tbl_pr.append(layout)
    total = int(sum(widths) * 1440)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")
    if indent.getparent() is None:
        tbl_pr.append(indent)
    grid_columns = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        width_dxa = int(width * 1440)
        if index < len(grid_columns):
            grid_columns[index].set(qn("w:w"), str(width_dxa))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(int(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_border(paragraph, side="bottom", color=NAVY, size=10, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
    if borders.getparent() is None:
        p_pr.append(borders)
    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    if border.getparent() is None:
        borders.append(border)


def shade_paragraph(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    p_pr.append(shading)


def add_page_field(paragraph, instruction):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instruction_run._r.append(instr)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    result_run = paragraph.add_run("1")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    for run in (begin_run, instruction_run, separate_run, result_run, end_run):
        set_run(run, "IBM Plex Mono", 8, MUTED)


def set_picture_alt(doc, alt):
    for doc_pr in doc.element.body.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt)
    for section in doc.sections:
        for doc_pr in section.header._element.xpath(".//wp:docPr"):
            doc_pr.set("descr", alt)
            doc_pr.set("title", alt)


def add_header_footer(doc, doc_type):
    section = doc.sections[0]
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    set_table_geometry(table, [1.1, 5.8])
    set_table_borders(table, WHITE, 0)
    mark_header_row(table.rows[0])
    logo_p = table.cell(0, 0).paragraphs[0]
    logo_p.add_run().add_picture(str(LOGO), width=Inches(0.48))
    title_p = table.cell(0, 1).paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(title_p.add_run(f"MASLOW AI · {doc_type}"), "IBM Plex Mono", 8, MUTED, True)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer_p.add_run("MASLOW BRAND OS 1.0.0 · "), "IBM Plex Mono", 8, MUTED)
    add_page_field(footer_p, "PAGE")
    set_run(footer_p.add_run(" / "), "IBM Plex Mono", 8, MUTED)
    add_page_field(footer_p, "NUMPAGES")


def base_document(title, subject, doc_type):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    define_styles(doc)
    add_header_footer(doc, doc_type)
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Maslow AI"
    props.keywords = "Maslow Brand OS 1.0.0, editable template, unresolved variables"
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    return doc


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1.25, 5.65])
    set_table_borders(table, LINE, 4)
    mark_header_row(table.rows[0])
    for row, (label, value) in zip(table.rows, rows):
        set_run(row.cells[0].paragraphs[0].add_run(label), "IBM Plex Mono", 8, PURPLE, True)
        set_run(row.cells[1].paragraphs[0].add_run(value), "Manrope", 10.5, TEXT)
    return table


def add_evidence_callout(doc, status="{{EVIDENCE_STATUS}}", source="{{SOURCE}}", text="{{CLAIM_OR_SCENARIO}}"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.9])
    set_table_borders(table, PURPLE, 6)
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, OFF_WHITE)
    label = cell.paragraphs[0]
    label.style = doc.styles["Maslow Evidence"]
    label.add_run(f"{status} · {source}")
    body = cell.add_paragraph(text)
    body.style = doc.styles["Normal"]


def add_title_block(doc, eyebrow, title, subtitle):
    p = doc.add_paragraph(eyebrow, style="Maslow Eyebrow")
    set_paragraph_border(p, "bottom", PURPLE, 8, 5)
    doc.add_paragraph(title, style="Maslow Title")
    doc.add_paragraph(subtitle, style="Maslow Subtitle")


def build_proposal():
    doc = base_document("Maslow proposal template", "Editable outcome-first proposal", "PROPOSAL")
    add_title_block(doc, "PROPOSAL · {{DATE}} · CONFIDENTIAL", "{{WAITING_WORKFLOW}}", "A working session for {{AUDIENCE}} to define ownership, the human decision, and the first deliverable.")
    add_metadata_table(doc, [("OWNER", "{{OWNER}}"), ("DECISION", "{{HUMAN_DECISION}}"), ("DELIVERABLE", "{{DELIVERABLE}}")])
    doc.add_paragraph("Scope", style="Heading 1")
    table = doc.add_table(rows=4, cols=3)
    set_table_geometry(table, [0.55, 2.1, 4.25])
    set_table_borders(table)
    headers = ["STEP", "PHASE", "OUTPUT"]
    for index, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], NAVY)
        set_run(table.rows[0].cells[index].paragraphs[0].add_run(value), "Manrope", 8.5, WHITE, True)
    mark_header_row(table.rows[0])
    for row_index in range(1, 4):
        values = [f"0{row_index}", f"{{{{PHASE_{row_index}}}}}", f"{{{{PHASE_{row_index}_OUTPUT}}}}"]
        for column, value in enumerate(values):
            set_run(table.rows[row_index].cells[column].paragraphs[0].add_run(value), "Manrope", 9.5, TEXT, column == 1)
    doc.add_paragraph("Evidence", style="Heading 1")
    add_evidence_callout(doc)
    action = doc.add_paragraph()
    action.paragraph_format.space_before = Pt(16)
    action.paragraph_format.space_after = Pt(0)
    shade_paragraph(action, NAVY)
    set_paragraph_border(action, "bottom", PINK, 18, 0)
    set_run(action.add_run("  {{CTA_LABEL}}  "), "Manrope", 10, WHITE, True)
    set_picture_alt(doc, "Maslow AI mark")
    doc.save(OUTPUT / "maslow-proposal-template.docx")


def build_memo():
    doc = base_document("Maslow memo template", "Editable decision memo", "INTERNAL MEMO")
    add_title_block(doc, "INTERNAL MEMO · {{DATE}}", "{{MEMO_TITLE}}", "{{ONE_SENTENCE_CONTEXT}}")
    add_metadata_table(doc, [("TO", "{{RECIPIENTS}}"), ("FROM", "{{AUTHOR}}"), ("RE", "{{SUBJECT}}"), ("STATUS", "{{DOCUMENT_STATUS}}")])
    doc.add_paragraph("Decision", style="Heading 1")
    doc.add_paragraph("{{DECISION}}")
    doc.add_paragraph("Owner and next action", style="Heading 1")
    doc.add_paragraph("{{OWNER_AND_ACTION}}")
    doc.add_paragraph("Evidence", style="Heading 1")
    add_evidence_callout(doc, text="{{EVIDENCE_NOTE}}")
    set_picture_alt(doc, "Maslow AI mark")
    doc.save(OUTPUT / "maslow-memo-template.docx")


def build_invoice():
    doc = base_document("Maslow invoice template", "Editable invoice", "INVOICE")
    add_title_block(doc, "INVOICE · {{INVOICE_NUMBER}}", "Invoice", "DUE {{DUE_DATE}}")
    meta = doc.add_table(rows=1, cols=3)
    set_table_geometry(meta, [3.0, 1.75, 2.15])
    set_table_borders(meta, LINE, 4)
    mark_header_row(meta.rows[0])
    values = [("BILL TO", "{{CLIENT_NAME}}\n{{CLIENT_CONTACT}}\n{{CLIENT_ADDRESS}}"), ("ISSUED", "{{ISSUE_DATE}}"), ("TERMS", "{{PAYMENT_TERMS}}")]
    for cell, (label, value) in zip(meta.rows[0].cells, values):
        set_run(cell.paragraphs[0].add_run(label), "IBM Plex Mono", 8, PURPLE, True)
        for line in value.split("\n"):
            paragraph = cell.add_paragraph()
            set_run(paragraph.add_run(line), "Manrope", 9.5, TEXT)
    doc.add_paragraph("Line items", style="Heading 1")
    table = doc.add_table(rows=4, cols=4)
    set_table_geometry(table, [3.65, 0.65, 1.25, 1.35])
    set_table_borders(table)
    for index, value in enumerate(["DESCRIPTION", "QTY", "RATE", "AMOUNT"]):
        set_cell_shading(table.rows[0].cells[index], NAVY)
        set_run(table.rows[0].cells[index].paragraphs[0].add_run(value), "Manrope", 8.5, WHITE, True)
    mark_header_row(table.rows[0])
    for row_index in range(1, 4):
        values = [f"{{{{ITEM_{row_index}}}}}", f"{{{{QTY_{row_index}}}}}", f"{{{{RATE_{row_index}}}}}", f"{{{{AMOUNT_{row_index}}}}}"]
        for column, value in enumerate(values):
            paragraph = table.rows[row_index].cells[column].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if column else WD_ALIGN_PARAGRAPH.LEFT
            set_run(paragraph.add_run(value), "Manrope", 9.5, TEXT)
    totals = doc.add_table(rows=3, cols=2)
    set_table_geometry(totals, [5.55, 1.35])
    set_table_borders(totals, LINE, 4)
    mark_header_row(totals.rows[0])
    for row, values in zip(totals.rows, [("SUBTOTAL", "{{SUBTOTAL}}"), ("TAX", "{{TAX}}"), ("TOTAL DUE", "{{TOTAL_DUE}}")]):
        if values[0] == "TOTAL DUE":
            set_cell_shading(row.cells[0], BAND)
            set_cell_shading(row.cells[1], BAND)
        color = WHITE if values[0] == "TOTAL DUE" else TEXT
        set_run(row.cells[0].paragraphs[0].add_run(values[0]), "Manrope", 9.5, color, True)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(row.cells[1].paragraphs[0].add_run(values[1]), "Manrope", 10.5, TEAL if values[0] == "TOTAL DUE" else color, True)
    doc.add_paragraph("Payment", style="Heading 1")
    doc.add_paragraph("{{PAYMENT_INSTRUCTIONS}}")
    contact = doc.add_paragraph()
    set_run(contact.add_run("{{CONTACT_LABEL}} >"), "Manrope", 9, LINK, True)
    set_picture_alt(doc, "Maslow AI mark")
    doc.save(OUTPUT / "maslow-invoice-template.docx")


def main():
    if not LOGO.exists():
        raise FileNotFoundError(f"Generate {LOGO} with scripts/build-social.mjs first")
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_proposal()
    build_memo()
    build_invoice()


if __name__ == "__main__":
    main()
